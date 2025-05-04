import streamlit as st
import openai
from openai import OpenAI
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from docx import Document
import tempfile
import os

# === STYLING ===
st.set_page_config(page_title="Book Summarizer", layout="centered")

# Inject custom CSS
st.markdown("""
    <style>
    body, div, input, label, textarea, select {
        font-family: Verdana !important;
        color: #2a2a2a !important;
    }
    h1 {
        font-family: 'Courier New', monospace !important;
        font-weight: normal !important;
        color: #2a2a2a !important;
        font-size: 40px !important;
        margin-bottom: 0.2em;
    }
    .by-chuck {
        font-size: 20px;
        font-weight: bold;
        letter-spacing: 2px;
        margin-bottom: 1em;
    }
    </style>
""", unsafe_allow_html=True)

# === HEADER ===
st.title("Book Summarizer")
st.markdown("""
<div class='by-chuck'>
    <span style='color:#f27802;'>C</span>
    <span style='color:#2e0854;'>H</span>
    <span style='color:#7786c8;'>U</span>
    <span style='color:#708090;'>C</span>
    <span style='color:#b02711;'>K</span>
</div>
""", unsafe_allow_html=True)

# === PASS PHRASE ===
passphrase = st.text_input("🔐 Enter passphrase to continue", type="password")
if passphrase.strip().lower() != "chucks books":
    st.stop()

# === FORM UI ===
book_title = st.text_input("Book Title")
book_notes = st.text_area("Notes, excerpts, or reflections (optional)")
summary_style = st.selectbox("Summary Style", ["Narrative", "Bullet", "Professional", "Reflective"])
submit = st.button("Generate Summary")

if submit and book_title:
    # === SETUP ===
    client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
    SCOPES = ['https://www.googleapis.com/auth/drive.file']
    SERVICE_ACCOUNT_INFO = st.secrets["gcp_service_account"]
    credentials = service_account.Credentials.from_service_account_info(SERVICE_ACCOUNT_INFO, scopes=SCOPES)
    drive_service = build('drive', 'v3', credentials=credentials)

    # === BUILD PROMPT ===
    if book_notes.strip():
        prompt = f"""
You are a literary critic and professional book summarizer.

Please provide a comprehensive summary, overview, and review of the book titled "{book_title}" using the following notes: {book_notes}.

Include:
- General summary of the book, its themes, and what makes it unique
- Thesis of the book and how it differs from other books in the same category or genre
- Main takeaways and actionable insights
- Chapter-by-chapter key ideas including quotes from each chapter and/or section
- Important quotes -- as many as possible (in-line and in a dedicated section)
- Distill a set of principles from the book that can be used in daily activities -- should be robust, 5-10 principles depending on the book 
- Include the book notes provided at the end of the this title "User Notes" 

Use this tone/style: {summary_style}.
Avoid markdown (**bold**, _italic_) — return clean plain text with clear structure.
        """
    else:
        prompt = f"""
You are a literary critic and professional book summarizer.

Please provide a comprehensive summary and overview of the book titled "{book_title}".

Include:
- General summary
- Thesis of the book
- Main takeaways
- Chapter-by-chapter key ideas
- Important quotes (in-line and in a dedicated section)

Use this tone/style: {summary_style}.
Avoid markdown (**bold**, _italic_) — return clean plain text with clear structure.
        """

    # === CALL OPENAI ===
    st.write("📤 Calling OpenAI...")
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a literary critic and professional book summarizer."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7
        )
        summary_text = response.choices[0].message.content.strip()
        st.success("✅ Summary generated!")
        st.text(f"Length: {len(summary_text)} characters")
    except Exception as e:
        st.error("❌ OpenAI API call failed")
        st.exception(e)
        st.stop()

    # === CREATE .DOCX ===
    try:
        st.write("📝 Creating Word document...")
        doc = Document()
        doc.add_heading(f"📘 {book_title}", level=1)

        for section in summary_text.split("\n\n"):
            lines = section.strip().split("\n")
            if len(lines) == 1:
                doc.add_paragraph(lines[0])
            else:
                doc.add_heading(lines[0], level=2)
                for line in lines[1:]:
                    if line.strip().startswith(("-", "•", "1.")):
                        doc.add_paragraph(line.strip()[2:], style='List Bullet')
                    else:
                        doc.add_paragraph(line.strip())

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

    except Exception as e:
        st.error("❌ Failed to generate Word document")
        st.exception(e)
        st.stop()

    # === UPLOAD TO GOOGLE DRIVE ===
    try:
        st.write("📤 Uploading to Google Drive...")

        file_metadata = {
            'name': f"Summary - {book_title}.docx",
            'parents': [st.secrets["GDRIVE_FOLDER_ID"]],
            'mimeType': 'application/vnd.google-apps.document'
        }

        media = MediaFileUpload(tmp_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id'
        ).execute()

        st.success("✅ Uploaded to Google Drive!")
        st.markdown(f"[📄 View in Drive](https://drive.google.com/file/d/{file['id']}/view)")

    except Exception as e:
        st.error("❌ Google Drive upload failed")
        st.exception(e)
        st.stop()

    os.remove(tmp_path)
